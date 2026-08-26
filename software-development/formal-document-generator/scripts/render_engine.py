#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""formal-document-generator 通用 docx 渲染引擎（内容与排版分离）

设计：
- 本文件不包含任何业务内容，只负责「排版」
- 业务内容由调用方提供 CONTENT 列表 + CONFIG 字典
- 调用方式：内容脚本 import build_document(CONFIG, CONTENT) 并执行

CONTENT 元素协议（元组）：
  ("h1", text) / ("h2", text) / ("h3", text)   —— 章节标题（Heading 样式，供目录域抓取）
  ("p", parts[, indent=True])                   —— 正文段落；parts 可为 str 或 [(text, bold), ...]
  ("b", parts)                                  —— 项目符号；parts 同上
  ("n", num, parts)                             —— 有序条目（num 为编号字符串或数字）
  ("callout", text[, title])                    —— 判断框（浅蓝底），title 默认「判断」
  ("warn", text[, title])                       —— 警示框（浅橙底），title 默认「待核实」
  ("table", headers, rows[, widths_cm][, note]) —— 表格；rows 元素可为 str 或 [(text,bold),...]；widths 为每列宽 cm 列表
  ("flow", text)                                —— 居中强调行（深蓝加粗）
  ("tag",)                                      —— 信息状态图例（四类标注说明）
  ("pb",)                                       —— 分页
  ("blank",)                                    —— 空段落

CONFIG 键（均可省略，有默认值）：
  TITLE / SUBTITLE / VERSION / DATE_LABEL / HEADER / FOOTER_NOTE / OUT / INTERNAL / TAG_LEGEND

用法示例见 templates/tech_transfer_discussion_template.py

已验证：磁电复合材料讨论稿 V0.1（14部分/27表/18项自检通过，2026-08-26）
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 色彩 ----------
DEEP_BLUE   = RGBColor(0x2B, 0x57, 0x9A)
GRAY_BLUE   = RGBColor(0x5A, 0x6E, 0x8C)
DARK_GRAY   = RGBColor(0x33, 0x33, 0x33)
RED         = RGBColor(0xC8, 0x00, 0x00)
DEEP_RED    = RGBColor(0xA0, 0x00, 0x00)
ORANGE      = RGBColor(0xC8, 0x78, 0x32)
LIGHT_BLUE  = "DCE6F1"   # 判断框浅蓝底
LIGHT_ORANGE= "FDE9D9"   # 待核实浅橙底
LIGHT_GRAY  = "F2F2F2"   # 隔行
WHITE       = "FFFFFF"

CJK_BODY = "PingFang SC"
CJK_HEAD = "PingFang SC"

DEFAULT_CONFIG = {
    "TITLE": "文档标题",
    "SUBTITLE": "",
    "VERSION": "V0.1",
    "DATE_LABEL": "",
    "HEADER": "",
    "FOOTER_NOTE": "",
    "OUT": "/Users/mac/Downloads/文档.docx",
    "INTERNAL": False,
    "TAG_LEGEND": "[已有材料] 项目方正式介绍；[公开可证] 权威公开渠道可验证；[待核实] 暂缺合同/报告/权属/客户证明；[方案假设] 为讨论方案而设的暂定条件。",
}


def set_run_font(run, size=10.5, bold=False, color=DARK_GRAY, name=CJK_BODY, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_field(run, instr):
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = instr
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def build_document(config, content):
    cfg = dict(DEFAULT_CONFIG)
    cfg.update(config or {})
    doc = Document()

    # ---------- 页面设置 ----------
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.2)
    sec.left_margin, sec.right_margin = Cm(2.3), Cm(2.3)

    # 页眉（无 HEADER 则不输出页眉文字，仅保留细线）
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cfg["HEADER"]:
        hr = hp.add_run(cfg["HEADER"])
        set_run_font(hr, 8.5, color=GRAY_BLUE)
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'B0B8C4')
    pBdr.append(bottom); pPr.append(pBdr)

    # 页脚（页码域）
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(cfg["FOOTER_NOTE"] + "第 ")
    set_run_font(fr, 8.5, color=GRAY_BLUE)
    r2 = fp.add_run(); set_run_font(r2, 8.5, color=GRAY_BLUE); add_field(r2, ' PAGE ')
    r3 = fp.add_run(" 页 / 共 "); set_run_font(r3, 8.5, color=GRAY_BLUE)
    r4 = fp.add_run(); set_run_font(r4, 8.5, color=GRAY_BLUE); add_field(r4, ' NUMPAGES ')
    r5 = fp.add_run(" 页"); set_run_font(r5, 8.5, color=GRAY_BLUE)

    # ---------- 样式 ----------
    normal = doc.styles['Normal']
    normal.font.name = CJK_BODY; normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), CJK_BODY)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.35

    for tag, size, color in [('Heading 1', 16, DEEP_BLUE), ('Heading 2', 13.5, GRAY_BLUE), ('Heading 3', 11.5, DARK_GRAY)]:
        st = doc.styles[tag]
        st.font.name = CJK_HEAD; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = color
        st.element.rPr.rFonts.set(qn('w:eastAsia'), CJK_HEAD)
        st.paragraph_format.space_before = Pt(14 if tag == 'Heading 1' else 10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    def H1(t): doc.add_heading(t, level=1)
    def H2(t): doc.add_heading(t, level=2)
    def H3(t): doc.add_heading(t, level=3)

    def P(parts, indent=True, size=10.5):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.35
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        if isinstance(parts, str):
            parts = [(parts, False)]
        for text, bold in parts:
            r = p.add_run(text)
            set_run_font(r, size, bold)
        return p

    def BULLET(parts, marker="\u2022"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(parts, str):
            parts = [(parts, False)]
        r = p.add_run(marker + " ")
        set_run_font(r, 10.5, True, GRAY_BLUE)
        for text, bold in parts:
            r = p.add_run(text)
            set_run_font(r, 10.5, bold)
        return p

    def NUM(n, parts):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(parts, str):
            parts = [(parts, False)]
        r = p.add_run(f"{n}. ")
        set_run_font(r, 10.5, True, GRAY_BLUE)
        for text, bold in parts:
            r = p.add_run(text)
            set_run_font(r, 10.5, bold)
        return p

    def CALLOUT(text, fill=LIGHT_BLUE, title="判断"):
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.cell(0, 0)
        shade_cell(cell, fill)
        p = cell.paragraphs[0]
        r = p.add_run(f"{title}：")
        set_run_font(r, 10.5, True, DEEP_BLUE if fill == LIGHT_BLUE else ORANGE)
        r2 = p.add_run(text)
        set_run_font(r2, 10.5, False, DARK_GRAY)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def TABLE(headers, rows, widths=None, size=9, note=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            shade_cell(hdr[i], "2B579A")
            p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h); set_run_font(r, size, True, RGBColor(255, 255, 255))
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci, val in enumerate(row):
                if ri % 2 == 1:
                    shade_cell(cells[ci], LIGHT_GRAY)
                p = cells[ci].paragraphs[0]
                parts = val if isinstance(val, list) else [(str(val), False)]
                for text, bold in parts:
                    r = p.add_run(text)
                    set_run_font(r, size, bold)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        if note:
            p = doc.add_paragraph()
            r = p.add_run(note)
            set_run_font(r, 8.5, False, GRAY_BLUE)
            p.paragraph_format.space_after = Pt(8)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def FLOW(text, size=10.5):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        set_run_font(r, size, True, DEEP_BLUE)
        return p

    def PAGEBREAK():
        doc.add_page_break()

    def TAG():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("信息状态说明：")
        set_run_font(r, 8.5, True, GRAY_BLUE)
        r = p.add_run(cfg["TAG_LEGEND"])
        set_run_font(r, 8.5, False, GRAY_BLUE)
        return p

    # ---------- 内部稿红标（CONFIG.INTERNAL=True 时） ----------
    if cfg.get("INTERNAL"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("\u3010机密 \u00b7 仅限董事长及核心决策层阅\u3011")
        set_run_font(r, 9, True, RED, name="Heiti SC")
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("\u2014 本件为内部决策参阅件，严禁外传 \u2014")
        set_run_font(r2, 9, False, DEEP_RED, name="Kaiti SC")
        PAGEBREAK()

    # ---------- 封面 ----------
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cfg["TITLE"])
    set_run_font(r, 22, True, DEEP_BLUE)
    if cfg["SUBTITLE"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cfg["SUBTITLE"])
        set_run_font(r, 12, False, GRAY_BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{cfg['VERSION']} · {cfg['DATE_LABEL']}")
    set_run_font(r, 11, True, DARK_GRAY)
    doc.add_paragraph()
    for line in ["内部讨论材料", "本文件用于讨论方向与机制，不构成投资承诺或最终交易方案。"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, 10, False, ORANGE if "不构成" in line else GRAY_BLUE)
    PAGEBREAK()

    # ---------- 目录（自动域） ----------
    p = doc.add_paragraph()
    r = p.add_run("目录")
    set_run_font(r, 16, True, DEEP_BLUE)
    toc_p = doc.add_paragraph()
    run = toc_p.add_run()
    set_run_font(run, 10.5)
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "（目录为自动域：打开文档后在目录上右键\u2192更新域\u2192更新整个目录，即可生成页码）"
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for el in (f1, it, f2, t, f3):
        run._r.append(el)
    PAGEBREAK()

    # ---------- 正文 ----------
    for item in content:
        kind = item[0]
        if kind == "h1": H1(item[1])
        elif kind == "h2": H2(item[1])
        elif kind == "h3": H3(item[1])
        elif kind == "p": P(item[1], indent=item[2] if len(item) > 2 else True)
        elif kind == "b": BULLET(item[1])
        elif kind == "n": NUM(item[1], item[2])
        elif kind == "callout": CALLOUT(item[1], fill=LIGHT_BLUE, title=item[2] if len(item) > 2 else "判断")
        elif kind == "warn": CALLOUT(item[1], fill=LIGHT_ORANGE, title=item[2] if len(item) > 2 else "待核实")
        elif kind == "table":
            TABLE(item[1], item[2], widths=item[3] if len(item) > 3 else None, note=item[4] if len(item) > 4 else None)
        elif kind == "flow": FLOW(item[1])
        elif kind == "pb": PAGEBREAK()
        elif kind == "tag": TAG()
        elif kind == "blank": doc.add_paragraph()
        else:
            print(f"[render_engine] 忽略未知元素: {kind}")

    doc.save(cfg["OUT"])
    print("saved:", cfg["OUT"])
    return cfg["OUT"]
