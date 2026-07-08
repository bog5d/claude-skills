#!/usr/bin/env python3
"""Reference pattern: generate structured Chinese document with python-docx + fpdf2.

This is a proven working example from the training plan generation session.
Key pattern: define content once as structured list, then feed to both renderers.

Pitfall fixed: fpdf2 multi_cell MUST use new_x="LMARGIN", new_y="NEXT" — otherwise
subsequent cell() calls fail with "Not enough horizontal space".
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from fpdf import FPDF

# ── Configuration ──
FONT_NAME = "PingFang SC"          # DOCX font
FONT_PATH = "/System/Library/Fonts/STHeiti Medium.ttc"  # PDF font (TTC)
OUT_DIR = os.path.expanduser("~/.hermes/cache/documents")

# ── Content definition: list of (type, data) ──
# type: h0/h1/h2/h3/p/ul/ol/table
# This single source feeds both DOCX and PDF renderers.

SECTIONS = []

def h0(t): SECTIONS.append(("h0", t))
def h1(t): SECTIONS.append(("h1", t))
def h2(t): SECTIONS.append(("h2", t))
def h3(t): SECTIONS.append(("h3", t))
def p(t):  SECTIONS.append(("p", t))
def ul(t): SECTIONS.append(("ul", t))
def ol(t): SECTIONS.append(("ol", t))
def tbl(data): SECTIONS.append(("table", data))

# ── DOCX Renderer ──
def set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_spacing(para, before=0, after=0, line_spacing=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def add_table_docx(doc, data):
    nrows, ncols = len(data), len(data[0])
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_run_font(run, size=10, bold=(i == 0))
            set_para_spacing(cell.paragraphs[0], before=2, after=2, line_spacing=1.1)
            if i == 0:  # header row
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): '2B579A', qn('w:val'): 'clear'})
                shading.append(shd)
                run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()

def build_docx(path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin = section.right_margin = Cm(2.8)

    for kind, data in SECTIONS:
        if kind == "h0":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(data)
            set_run_font(run, size=22, bold=True)
            set_para_spacing(p, before=0, after=12)
        elif kind == "h1":
            p = doc.add_paragraph()
            run = p.add_run(data)
            set_run_font(run, size=15, bold=True, color=(43, 87, 154))
            set_para_spacing(p, before=18, after=6)
        elif kind == "h2":
            p = doc.add_paragraph()
            run = p.add_run(data)
            set_run_font(run, size=13, bold=True, color=(43, 87, 154))
            set_para_spacing(p, before=12, after=4)
        elif kind == "h3":
            p = doc.add_paragraph()
            run = p.add_run(data)
            set_run_font(run, size=11.5, bold=True)
            set_para_spacing(p, before=10, after=3)
        elif kind == "p":
            if not data.strip(): continue
            p = doc.add_paragraph()
            run = p.add_run(data)
            set_run_font(run, size=11)
            set_para_spacing(p, before=0, after=6, line_spacing=1.5)
        elif kind == "ul":
            for li in data:
                p = doc.add_paragraph(style='List Bullet')
                p.clear()
                run = p.add_run(li)
                set_run_font(run, size=11)
                set_para_spacing(p, before=0, after=2, line_spacing=1.3)
        elif kind == "ol":
            for li in data:
                p = doc.add_paragraph(style='List Number')
                p.clear()
                run = p.add_run(li)
                set_run_font(run, size=11)
                set_para_spacing(p, before=0, after=2, line_spacing=1.3)
        elif kind == "table":
            add_table_docx(doc, data)

    doc.save(path)
    print(f"DOCX: {path}")

# ── PDF Renderer (fpdf2) ──
class StructuredPDF(FPDF):
    def __init__(self, font_path):
        super().__init__('P', 'mm', 'A4')
        self.add_font("zh", "", font_path)
        self.add_font("zh", "B", font_path)
        self.set_auto_page_break(True, 20)
        self.font_path = font_path

    def header(self):
        if self.page_no() > 1:
            self.set_font("zh", "", 7)
            self.set_text_color(150, 150, 150)
            self.cell(0, 5, "Document Title", align='R')
            self.ln(8)

    def footer(self):
        self.set_y(-15)
        self.set_font("zh", "", 7)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}", align='C')

    def check_break(self, needed=30):
        if self.get_y() > self.h - self.b_margin - needed:
            self.add_page()

    def write_ul(self, items):
        self.set_font("zh", "", 10.5)
        self.set_text_color(40, 40, 40)
        for li in items:
            self.cell(6, 6, "\u2022")  # bullet
            # CRITICAL: new_x must be "LMARGIN"
            self.multi_cell(0, 6, li, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def write_ol(self, items):
        self.set_font("zh", "", 10.5)
        self.set_text_color(40, 40, 40)
        for i, li in enumerate(items, 1):
            self.cell(8, 6, f"{i}.")
            # CRITICAL: new_x must be "LMARGIN"
            self.multi_cell(0, 6, li, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def write_table(self, data):
        ncols = len(data[0])
        usable = self.w - self.l_margin - self.r_margin
        widths = [usable / ncols] * ncols
        # Header
        self.set_fill_color(43, 87, 154)
        self.set_text_color(255, 255, 255)
        self.set_font("zh", "B", 9.5)
        for j, cell in enumerate(data[0]):
            self.cell(widths[j], 8, cell, border=1, fill=True, align='C')
        self.ln()
        # Data rows (alternating background)
        self.set_font("zh", "", 9.5)
        for i, row in enumerate(data[1:]):
            self.set_fill_color(240, 245, 252) if i % 2 else self.set_fill_color(255, 255, 255)
            self.set_text_color(40, 40, 40)
            for j, cell in enumerate(row):
                self.cell(widths[j], 8, str(cell), border=1, fill=True, align='L' if j == 0 else 'C')
            self.ln()
        self.ln(4)

def build_pdf(path, font_path):
    pdf = StructuredPDF(font_path)
    pdf.add_page()
    for kind, data in SECTIONS:
        if kind == "h0": continue
        elif kind == "h1":
            pdf.check_break(20)
            pdf.set_font("zh", "B", 14)
            pdf.set_text_color(43, 87, 154)
            pdf.cell(0, 10, data); pdf.ln(12)
        elif kind == "h2":
            pdf.check_break(20)
            pdf.set_font("zh", "B", 12)
            pdf.set_text_color(43, 87, 154)
            pdf.cell(0, 8, data); pdf.ln(10)
        elif kind == "h3":
            pdf.check_break(20)
            pdf.set_font("zh", "B", 11)
            pdf.set_text_color(60, 60, 60)
            pdf.cell(0, 7, data); pdf.ln(9)
        elif kind == "p":
            if data.strip():
                pdf.set_font("zh", "", 10.5)
                pdf.set_text_color(40, 40, 40)
                pdf.multi_cell(0, 6.5, data)
                pdf.ln(2)
        elif kind == "ul":
            pdf.write_ul(data)
        elif kind == "ol":
            pdf.write_ol(data)
        elif kind == "table":
            est_h = len(data) * 8 + 15
            pdf.check_break(est_h)
            pdf.write_table(data)
    pdf.output(path)
    print(f"PDF: {pdf_path}")

# ── Usage ──
if __name__ == "__main__":
    # Define content here...
    h0("Document Title")
    p("Body text goes here.")
    docx_path = os.path.join(OUT_DIR, "output.docx")
    pdf_path = os.path.join(OUT_DIR, "output.pdf")
    build_docx(docx_path)
    build_pdf(pdf_path, FONT_PATH)
    print("Done.")
