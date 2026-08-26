#!/usr/bin/env python3
"""CONTENT 驱动 docx 渲染器模板（2026-08-26 实测）。

用法：
1. 写内容文件 xxx_content.py（辅助函数 h1/h2/h3/p/b/n/callout/warn/table/flow/pb/tag 追加到 CONTENT 列表）
2. 把本文件复制为 /tmp/xxx_render.py，将 import 改为 from xxx_content import CONTENT
3. python3 xxx_render.py → 输出 docx（含封面/目录域/页眉页脚/页码/底纹/判断框）
"""
import sys
sys.path.insert(0, '/tmp')
from discussion_content import CONTENT  # 改成你的内容模块

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 色彩 ----------
DEEP_BLUE   = RGBColor(0x2B, 0x57, 0x9A)
GRAY_BLUE   = RGBColor(0x5A, 0x6E, 0x8C)
DARK_GRAY   = RGBColor(0x33, 0x33, 0x33)
ORANGE      = RGBColor(0xC8, 0x78, 0x32)
LIGHT_BLUE  = "DCE6F1"   # 判断框浅蓝底
LIGHT_ORANGE= "FDE9D9"   # 待核实浅橙底
LIGHT_GRAY  = "F2F2F2"   # 隔行
CJK_BODY = "PingFang SC"

def set_run_font(run, size=10.5, bold=False, color=DARK_GRAY, name=CJK_BODY):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = name; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_field(run, instr):
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.2)
sec.left_margin = sec.right_margin = Cm(2.3)

# 页眉（带下边框线）
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(hp.add_run("文档标题（替换）"), 8.5, color=GRAY_BLUE)
pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'4'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'B0B8C4')
pBdr.append(bottom); hp._p.get_or_add_pPr().append(pBdr)

# 页脚（版本 + 页码域）
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(fp.add_run("V0.1 · "), 8.5, color=GRAY_BLUE)
set_run_font(fp.add_run("第 "), 8.5, color=GRAY_BLUE)
add_field(fp.add_run(), ' PAGE ')
set_run_font(fp.add_run(" 页 / 共 "), 8.5, color=GRAY_BLUE)
add_field(fp.add_run(), ' NUMPAGES ')
set_run_font(fp.add_run(" 页"), 8.5, color=GRAY_BLUE)

# 样式：Normal + Heading（TOC 依赖内置 Heading，需覆盖字体）
normal = doc.styles['Normal']
normal.font.name = CJK_BODY; normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), CJK_BODY)
normal.paragraph_format.line_spacing = 1.35
for tag, size, color in [('Heading 1', 16, DEEP_BLUE), ('Heading 2', 13.5, GRAY_BLUE), ('Heading 3', 11.5, DARK_GRAY)]:
    st = doc.styles[tag]
    st.font.name = CJK_BODY; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = color
    st.element.rPr.rFonts.set(qn('w:eastAsia'), CJK_BODY)
    st.paragraph_format.space_before = Pt(14 if tag == 'Heading 1' else 10)
    st.paragraph_format.space_after = Pt(6)

def P(parts, size=10.5):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.first_line_indent = Cm(0.74)
    if isinstance(parts, str): parts = [(parts, False)]
    for text, bold in parts:
        set_run_font(p.add_run(text), size, bold)
    return p

def BULLET(parts, marker="\u2022"):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6); p.paragraph_format.space_after = Pt(2)
    if isinstance(parts, str): parts = [(parts, False)]
    set_run_font(p.add_run(marker + " "), 10.5, True, GRAY_BLUE)
    for text, bold in parts:
        set_run_font(p.add_run(text), 10.5, bold)
    return p

def CALLOUT(text, fill=LIGHT_BLUE, title="判断"):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0); shade_cell(cell, fill)
    set_run_font(cell.paragraphs[0].add_run(f"{title}："), 10.5, True, DEEP_BLUE if fill == LIGHT_BLUE else ORANGE)
    set_run_font(cell.paragraphs[0].add_run(text), 10.5, False, DARK_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def TABLE(headers, rows, widths=None, size=9, note=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade_cell(c, "2B579A")
        set_run_font(c.paragraphs[0].add_run(h), size, True, RGBColor(255,255,255))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1: shade_cell(cells[ci], LIGHT_GRAY)
            parts = val if isinstance(val, list) else [(str(val), False)]
            for text, bold in parts:
                set_run_font(cells[ci].paragraphs[0].add_run(text), size, bold)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)  # 必须逐行设置
    if note:
        p = doc.add_paragraph(); set_run_font(p.add_run(note), 8.5, False, GRAY_BLUE)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

def FLOW(text, size=10.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), size, True, DEEP_BLUE)
    return p

def PAGEBREAK(): doc.add_page_break()

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run(); set_run_font(run, 10.5)
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "（自动域：右键→更新域→更新整个目录）"
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for el in (f1, it, f2, t, f3): run._r.append(el)

# ========== 封面（按需替换） ==========
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("主标题（替换）"), 22, True, DEEP_BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("——副标题——"), 12, False, GRAY_BLUE)
PAGEBREAK()
add_toc()
PAGEBREAK()

# ========== 正文（遍历 CONTENT） ==========
for item in CONTENT:
    kind = item[0]
    if kind == "h1": doc.add_heading(item[1], level=1)
    elif kind == "h2": doc.add_heading(item[1], level=2)
    elif kind == "h3": doc.add_heading(item[1], level=3)
    elif kind == "p": P(item[1])
    elif kind == "b": BULLET(item[1])
    elif kind == "n":
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
        set_run_font(p.add_run(f"{item[1]}. "), 10.5, True, GRAY_BLUE)
        if isinstance(item[2], str): item[2] = [(item[2], False)]
        for text, bold in item[2]: set_run_font(p.add_run(text), 10.5, bold)
    elif kind == "callout": CALLOUT(item[1], title=item[2] if len(item) > 2 else "判断")
    elif kind == "warn": CALLOUT(item[1], fill=LIGHT_ORANGE, title=item[2] if len(item) > 2 else "待核实")
    elif kind == "table": TABLE(item[1], item[2], widths=item[3] if len(item) > 3 else None, note=item[4] if len(item) > 4 else None)
    elif kind == "flow": FLOW(item[1])
    elif kind == "pb": PAGEBREAK()
    elif kind == "blank": doc.add_paragraph()

OUT = "/Users/mac/Downloads/输出文档.docx"
doc.save(OUT)
print("saved:", OUT)
