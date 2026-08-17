# -*- coding: utf-8 -*-
"""Word + PDF 双份中文汇报生成器（2026-08-17 验证可跑，macOS）
用法：改 TITLE/SUBTITLE/SECTIONS 后直接 python3 运行。
输出：~/.hermes/cache/documents/<主题>/ 下 .docx + .pdf（Telegram MEDIA 白名单）。
"""
import os

OUT_DIR = os.path.expanduser("~/.hermes/cache/documents/汇报材料/")
os.makedirs(OUT_DIR, exist_ok=True)
DOCX_PATH = os.path.join(OUT_DIR, "汇报_YYYYMMDD.docx")
PDF_PATH = os.path.join(OUT_DIR, "汇报_YYYYMMDD.pdf")

TITLE = "汇报标题"
SUBTITLE = "副标题 ｜ 日期"

# 每节：(章节标题, [(小标题, 正文), ...])
SECTIONS = [
    ("一、示例章节", [
        ("要点A", "正文内容……"),
        ("要点B", "正文内容……"),
    ]),
]

# ============ Word ============
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Helvetica"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")  # 关键：中文字体

def set_cn(run, name="STHeiti"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(TITLE); r.font.size = Pt(20); r.font.bold = True; set_cn(r)
r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x66)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(SUBTITLE); r.font.size = Pt(11); set_cn(r)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

for sec_title, items in SECTIONS:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(sec_title); r.font.size = Pt(14); r.font.bold = True; set_cn(r)
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x66)
    for head, body in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(f"▍{head}：")
        r.font.bold = True; r.font.size = Pt(10.5); set_cn(r)
        r.font.color.rgb = RGBColor(0x8A, 0x4B, 0x00)
        r2 = p.add_run(body); r2.font.size = Pt(10.5); set_cn(r2)

doc.save(DOCX_PATH)

# ============ PDF ============
from fpdf import FPDF

class PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font("STHeiti", "", "/System/Library/Fonts/STHeiti Medium.ttc")
        self.add_font("STHeiti", "B", "/System/Library/Fonts/STHeiti Medium.ttc")

pdf = PDF()
pdf.set_auto_page_break(auto=True, margin=18)
pdf.add_page(); pdf.set_margins(16, 16, 16)

pdf.set_font("STHeiti", "B", 18)
pdf.set_text_color(0x1F, 0x3B, 0x66)
pdf.cell(0, 12, TITLE, new_x="LMARGIN", new_y="NEXT", align="C")
pdf.set_font("STHeiti", "", 10)
pdf.set_text_color(0x88, 0x88, 0x88)
pdf.cell(0, 7, SUBTITLE, new_x="LMARGIN", new_y="NEXT", align="C")
pdf.ln(4)

for sec_title, items in SECTIONS:
    pdf.set_font("STHeiti", "B", 13)
    pdf.set_text_color(0x1F, 0x3B, 0x66)
    pdf.multi_cell(0, 8, sec_title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)
    for head, body in items:
        pdf.set_font("STHeiti", "B", 10)
        pdf.set_text_color(0x8A, 0x4B, 0x00)
        pdf.multi_cell(0, 6, f"▍{head}：", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("STHeiti", "", 10)
        pdf.set_text_color(0x22, 0x22, 0x22)
        pdf.multi_cell(0, 6, body, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
    pdf.ln(2)

pdf.output(PDF_PATH)
print(f"DOCX OK: {DOCX_PATH}")
print(f"PDF OK: {PDF_PATH} ({os.path.getsize(PDF_PATH)//1024}KB)")

# 交付前：ls -la 确认双份在位 → MEDIA: 双发（勿先问格式）
