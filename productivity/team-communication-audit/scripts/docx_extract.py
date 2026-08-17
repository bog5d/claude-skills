# -*- coding: utf-8 -*-
"""docx -> txt 批量转换 + 提取分析文件 4/5/6 节（口径偏差/团队盲区/承诺与风险）
用法1（转换）：python3 docx_extract.py --convert <raw_dir> <out_dir>
用法2（提取节）：python3 docx_extract.py --extract <analysis_dir> <out_dir>
注：用 terminal + python3 跑；execute_code 可能被环境拦。
"""
import os, sys, glob, re


def convert_docx_dir(raw_dir: str, out_dir: str):
    from docx import Document
    os.makedirs(out_dir, exist_ok=True)
    files = sorted(glob.glob(os.path.join(raw_dir, "**", "*.docx"), recursive=True))
    for f in files:
        doc = Document(f)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        base = os.path.splitext(os.path.basename(f))[0]
        out = os.path.join(out_dir, base + ".txt")
        with open(out, "w", encoding="utf-8") as fh:
            fh.write("\n".join(paras))
        print(f"{base}.txt  {os.path.getsize(out)//1024}KB  {len(paras)}段")
    print(f"共 {len(files)} 份 -> {out_dir}")


def extract_sections(analysis_dir: str, out_dir: str):
    """从分场分析 md 中切出 4/5/6/7 节（口径偏差/团队盲区/承诺与风险/待确认）"""
    os.makedirs(out_dir, exist_ok=True)
    files = sorted(glob.glob(os.path.join(analysis_dir, "*.分析.md")))
    for f in files:
        with open(f, encoding="utf-8") as fh:
            text = fh.read()
        parts = []
        for sec_no, sec_name in [(4, "口径偏差"), (5, "团队盲区"), (6, "承诺与风险"), (7, "待")]:
            m = re.search(rf"## {sec_no}\. {sec_name}(.*?)(?=\n## |\Z)", text, re.S)
            if m:
                parts.append(f"## {sec_no}. {sec_name}\n{m.group(1).strip()}")
        base = os.path.splitext(os.path.basename(f))[0]
        out = os.path.join(out_dir, base + "_4567.md")
        with open(out, "w", encoding="utf-8") as fh:
            fh.write(f"# {base}\n\n" + "\n\n".join(parts))
        print(f"{base}: {os.path.getsize(out)//1024}KB")


if __name__ == "__main__":
    mode = sys.argv[1] if len(sys.argv) > 1 else "convert"
    if mode == "--convert" and len(sys.argv) >= 4:
        convert_docx_dir(sys.argv[2], sys.argv[3])
    elif mode == "--extract" and len(sys.argv) >= 4:
        extract_sections(sys.argv[2], sys.argv[3])
    else:
        print(__doc__)
